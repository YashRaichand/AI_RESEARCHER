import re
import fitz  # PyMuPDF
import pdfplumber
from pathlib import Path
from typing import Dict, Any, List
import docx


class PDFEngine:
    def __init__(self):
        self.equation_pattern = re.compile(
            r'(\$[^$]+\$|\\\([^)]+\\\)|\\\[[^\]]+\\\]|\\begin\{equation\}.*?\\end\{equation\})',
            re.DOTALL
        )

    def extract(self, file_path: str, file_type: str) -> Dict[str, Any]:
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        else:
            return self._extract_txt(file_path)

    def _extract_pdf(self, path: str) -> Dict[str, Any]:
        result = {
            "title": "", "authors": [], "abstract": "", "keywords": [],
            "sections": {}, "references": [], "figures": [], "tables": [],
            "equations": [], "full_text": "", "page_count": 0, "word_count": 0,
        }
        try:
            doc = fitz.open(path)
            result["page_count"] = len(doc)
            full_text_parts = []
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                full_text_parts.append(text)
                images = page.get_images()
                for img_idx, img in enumerate(images):
                    result["figures"].append({
                        "page": page_num + 1,
                        "index": img_idx,
                        "caption": f"Figure on page {page_num + 1}",
                    })
            doc.close()
            full_text = "\n".join(full_text_parts)
            result["full_text"] = full_text
            result["word_count"] = len(full_text.split())
            result["equations"] = self.detect_equations(full_text)
            result["title"] = self._extract_title(full_text)
            result["authors"] = self._extract_authors(full_text)
            result["abstract"] = self._extract_abstract(full_text)
            result["keywords"] = self._extract_keywords(full_text)
            result["references"] = self._extract_references(full_text)
        except Exception as e:
            result["error"] = str(e)

        try:
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for tbl_idx, table in enumerate(tables):
                        if table:
                            result["tables"].append({
                                "page": page_num + 1,
                                "index": tbl_idx,
                                "data": table,
                            })
        except Exception:
            pass

        return result

    def _extract_docx(self, path: str) -> Dict[str, Any]:
        result = {
            "title": "", "authors": [], "abstract": "", "keywords": [],
            "sections": {}, "references": [], "figures": [], "tables": [],
            "equations": [], "full_text": "", "page_count": 1, "word_count": 0,
        }
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        result["full_text"] = full_text
        result["word_count"] = len(full_text.split())
        result["title"] = paragraphs[0] if paragraphs else "Untitled"
        result["equations"] = self.detect_equations(full_text)
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            result["tables"].append({"page": 1, "data": rows})
        return result

    def _extract_txt(self, path: str) -> Dict[str, Any]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
        return {
            "title": full_text.split("\n")[0][:200] if full_text else "Untitled",
            "authors": [], "abstract": "", "keywords": [],
            "sections": {}, "references": [], "figures": [], "tables": [],
            "equations": self.detect_equations(full_text),
            "full_text": full_text,
            "page_count": 1,
            "word_count": len(full_text.split()),
        }

    def detect_equations(self, text: str) -> List[str]:
        equations = self.equation_pattern.findall(text)
        math_patterns = re.findall(r'[A-Za-z]\s*=\s*[^,\.\n]{3,50}', text)
        return list(set(equations + math_patterns))[:50]

    def _extract_title(self, text: str) -> str:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:10]:
            if 10 < len(line) < 300 and not line.startswith("http"):
                return line
        return lines[0][:200] if lines else "Untitled"

    def _extract_authors(self, text: str) -> List[str]:
        author_pattern = re.compile(
            r'(?:Authors?|By)[:\s]+([A-Z][a-z]+ [A-Z][a-z]+(?:,\s*[A-Z][a-z]+ [A-Z][a-z]+)*)',
            re.IGNORECASE
        )
        match = author_pattern.search(text[:2000])
        if match:
            return [a.strip() for a in match.group(1).split(",")]
        names = re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', text[:1000])
        return list(set(names))[:5]

    def _extract_abstract(self, text: str) -> str:
        abstract_match = re.search(
            r'(?:Abstract|ABSTRACT)[:\s]*\n?(.*?)(?:\n\n|\n(?:Introduction|1\.|Keywords))',
            text, re.DOTALL | re.IGNORECASE
        )
        if abstract_match:
            return abstract_match.group(1).strip()[:2000]
        return ""

    def _extract_keywords(self, text: str) -> List[str]:
        kw_match = re.search(
            r'(?:Keywords?|Index Terms?)[:\s]+(.*?)(?:\n\n|\n[A-Z])',
            text, re.IGNORECASE | re.DOTALL
        )
        if kw_match:
            raw = kw_match.group(1)
            keywords = re.split(r'[,;·•]', raw)
            return [k.strip() for k in keywords if 2 < len(k.strip()) < 50][:20]
        return []

    def _extract_references(self, text: str) -> List[str]:
        ref_match = re.search(
            r'(?:References|Bibliography|REFERENCES)\s*\n(.*?)$',
            text, re.DOTALL | re.IGNORECASE
        )
        if ref_match:
            refs_text = ref_match.group(1)
            refs = re.split(r'\n\[?\d+\]?\.?\s+', refs_text)
            return [r.strip() for r in refs if len(r.strip()) > 20][:100]
        return []
